"""Word (.docx) report builder using python-docx — a branded, client-ready audit document."""
from __future__ import annotations

import io

from ..config import get_settings
from ..schemas import AuditResult, Priority, Severity

_SEV_HEX = {
    Severity.critical: "C0392B", Severity.high: "E67E22", Severity.medium: "D4AC0D",
    Severity.low: "2E86C1", Severity.info: "7F8C8D",
}


def _hex(s: str) -> "RGBColor":  # type: ignore[name-defined]
    from docx.shared import RGBColor
    return RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))


def docx_available() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except Exception:
        return False


def render_docx(audit: AuditResult) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    cfg = get_settings()
    analysis = audit.analysis
    crawl = audit.crawl
    client = audit.client_name or (crawl.domain if crawl else audit.url)
    primary = cfg.agency_primary_color.lstrip("#") or "2563EB"
    score = audit.score if audit.score is not None else (analysis.score if analysis else 0)

    doc = Document()
    doc.core_properties.title = f"Website Audit — {client}"
    doc.core_properties.author = cfg.agency_name

    # cover
    brand = doc.add_paragraph()
    run = brand.add_run(f"🩺 {cfg.product_name}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = _hex(primary)

    title = doc.add_heading("Website Audit Report", level=0)
    doc.add_paragraph(f"Prepared for {client} by {cfg.agency_name}")
    meta = doc.add_paragraph()
    meta.add_run(
        f"Website: {audit.url}    |    Pages: {len(crawl.pages) if crawl else 0}    |    "
        f"Links checked: {crawl.links_checked if crawl else 0}    |    "
        f"HTTPS: {'Yes' if crawl and crawl.https else 'No'}"
    ).italic = True

    # overall score
    p = doc.add_paragraph()
    r = p.add_run(f"Overall website health: {score}/100")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = _hex("16A34A" if score >= 80 else "D97706" if score >= 50 else "C0392B")

    if analysis:
        doc.add_heading("Executive summary", level=1)
        doc.add_paragraph(analysis.executive_summary)

        # category scores table
        if analysis.category_scores:
            doc.add_heading("Category scores", level=1)
            t = doc.add_table(rows=1, cols=3)
            t.style = "Light Grid Accent 1"
            h = t.rows[0].cells
            h[0].text, h[1].text, h[2].text = "Category", "Score", "Issues"
            for c in analysis.category_scores:
                row = t.add_row().cells
                row[0].text = c.category.value
                row[1].text = f"{c.score}/100"
                row[2].text = str(c.issue_count)

        # core web vitals
        v = crawl.vitals if crawl else None
        if v and v.source == "psi":
            doc.add_heading("Core Web Vitals (mobile)", level=1)
            doc.add_paragraph(
                f"Performance: {v.performance_score or '—'} | "
                f"LCP: {f'{v.lcp_ms/1000:.1f}s' if v.lcp_ms else '—'} | "
                f"CLS: {f'{v.cls:.3f}' if v.cls is not None else '—'} | "
                f"INP: {f'{v.inp_ms:.0f}ms' if v.inp_ms else '—'}"
            )

        # quick wins
        if analysis.quick_wins:
            doc.add_heading("Quick wins", level=1)
            for w in analysis.quick_wins:
                doc.add_paragraph(w, style="List Bullet")

        # recommended fixes table
        doc.add_heading("Recommended fixes", level=1)
        ft = doc.add_table(rows=1, cols=4)
        ft.style = "Light Grid Accent 1"
        hc = ft.rows[0].cells
        hc[0].text, hc[1].text, hc[2].text, hc[3].text = "Issue", "Category", "Impact", "Difficulty"
        impact_rank = {"High": 0, "Medium": 1, "Low": 2}
        fixes = sorted(analysis.issues, key=lambda i: impact_rank.get(i.impact.value, 3))[:15]
        for i in fixes:
            row = ft.add_row().cells
            row[0].text, row[1].text = i.title, i.category.value
            row[2].text, row[3].text = i.impact.value, i.difficulty.value

        # action plan
        doc.add_heading("Action plan", level=1)
        for prio in Priority:
            items = [i for i in analysis.issues if i.priority == prio]
            if not items:
                continue
            doc.add_heading(f"{prio.value} ({len(items)})", level=2)
            for i in items:
                hp = doc.add_paragraph()
                hr = hp.add_run(f"{i.title}  ·  {i.category.value}")
                hr.bold = True
                hr.font.color.rgb = _hex(_SEV_HEX.get(i.severity, "333333"))
                doc.add_paragraph(i.description)
                fp = doc.add_paragraph()
                fp.add_run("Fix: ").bold = True
                fp.add_run(f"{i.recommendation}  ({i.impact.value} impact · {i.difficulty.value})")
                if i.affected:
                    a = doc.add_paragraph(f"Affected: {i.affected}")
                    a.runs[0].italic = True

        if analysis.expected_results:
            doc.add_heading("Expected results", level=1)
            for r2 in analysis.expected_results:
                doc.add_paragraph(r2, style="List Bullet")

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = foot.add_run(
        f"Generated by {cfg.product_name} · {cfg.agency_name} · {cfg.agency_website} · {cfg.agency_email}"
    )
    fr.italic = True
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x90, 0x94, 0x9C)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
